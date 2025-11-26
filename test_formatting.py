"""
Test script for add_paragraph and add_heading formatting parameters.
"""
import asyncio
from docx import Document
from word_document_server.tools.content_tools import add_paragraph, add_heading
from word_document_server.tools.document_tools import create_document


async def test_formatting():
    """Test the new formatting parameters."""
    test_doc = 'test_formatting.docx'

    # Create test document
    print("Creating test document...")
    await create_document(test_doc, title="Formatting Test", author="Test Script")

    # Test 1: Name with large font
    print("Test 1: Adding name with large Helvetica 36pt bold...")
    result = await add_paragraph(
        test_doc,
        "JAMES MEHORTER",
        font_name="Helvetica",
        font_size=36,
        bold=True
    )
    print(f"  Result: {result}")

    # Test 2: Title line
    print("Test 2: Adding title with Helvetica 14pt...")
    result = await add_paragraph(
        test_doc,
        "Principal Software Engineer | Technical Team Lead",
        font_name="Helvetica",
        font_size=14
    )
    print(f"  Result: {result}")

    # Test 3: Section header with border
    print("Test 3: Adding section header with border...")
    result = await add_heading(
        test_doc,
        "PROFESSIONAL SUMMARY",
        level=2,
        font_name="Helvetica",
        font_size=14,
        bold=True,
        border_bottom=True
    )
    print(f"  Result: {result}")

    # Test 4: Body text in Times New Roman
    print("Test 4: Adding body text in Times New Roman 14pt...")
    result = await add_paragraph(
        test_doc,
        "This is body text that should be in Times New Roman at 14pt. "
        "It demonstrates the ability to apply different fonts to different paragraphs.",
        font_name="Times New Roman",
        font_size=14
    )
    print(f"  Result: {result}")

    # Test 5: Another section header
    print("Test 5: Adding another section header with border...")
    result = await add_heading(
        test_doc,
        "SKILLS",
        level=2,
        font_name="Helvetica",
        font_size=14,
        bold=True,
        border_bottom=True
    )
    print(f"  Result: {result}")

    # Test 6: Italic text with color
    print("Test 6: Adding italic text with color...")
    result = await add_paragraph(
        test_doc,
        "This text is italic and colored blue.",
        font_name="Arial",
        font_size=12,
        italic=True,
        color="0000FF"
    )
    print(f"  Result: {result}")

    print(f"\n✅ Test document created: {test_doc}")

    # Verify formatting
    print("\nVerifying formatting...")
    verify_doc = Document(test_doc)
    for i, para in enumerate(verify_doc.paragraphs):
        if para.runs:
            run = para.runs[0]
            text_preview = para.text[:50] + "..." if len(para.text) > 50 else para.text
            print(f"\nParagraph {i}: {text_preview}")
            print(f"  Font: {run.font.name}")
            print(f"  Size: {run.font.size}")
            print(f"  Bold: {run.font.bold}")
            print(f"  Italic: {run.font.italic}")

    print("\n✅ All tests completed successfully!")
    print(f"Open {test_doc} in Word to verify the formatting visually.")


if __name__ == "__main__":
    asyncio.run(test_formatting())
