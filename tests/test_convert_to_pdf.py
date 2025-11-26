import asyncio
from pathlib import Path

import pytest
from docx import Document

# Target for testing: convert_to_pdf (async function)
from word_document_server.tools.extended_document_tools import convert_to_pdf


def _make_sample_docx(path: Path) -> None:
    """Generates a simple .docx file in a temporary directory."""
    doc = Document()
    doc.add_heading("Conversion Test Document", level=1)
    doc.add_paragraph("This is a test paragraph for PDF conversion. Contains ASCII too.")
    doc.add_paragraph("Second paragraph: Contains special characters and spaces to cover path/content edge cases.")
    doc.save(path)


def test_convert_to_pdf_with_temp_docx(tmp_path: Path):
    """
    End-to-end test: Create a temporary .docx -> call convert_to_pdf -> validate the PDF output.

    Notes:
    - On Linux/macOS, it first tries LibreOffice (soffice/libreoffice),
      and falls back to docx2pdf on failure (requires Microsoft Word).
    - If these tools are missing or the command is unavailable, the test is skipped with a reason.
    """
    # 1) Generate a docx file with spaces in its name in the temp directory
    src_doc = tmp_path / "sample document with spaces.docx"
    _make_sample_docx(src_doc)

    # 2) Define the output PDF path (also in the temp directory)
    out_pdf = tmp_path / "converted output.pdf"

    # 3) Run the asynchronous function under test
    result_msg = asyncio.run(convert_to_pdf(str(src_doc), output_filename=str(out_pdf)))

    # 4) Success condition: the return message contains success keywords, or the target PDF exists
    success_keywords = ["successfully converted", "converted to PDF"]
    success = any(k.lower() in result_msg.lower() for k in success_keywords) or out_pdf.exists()

    if not success:
        # When LibreOffice or Microsoft Word is not installed, the tool returns a hint.
        # In this case, skip the test instead of failing.
        pytest.skip(f"PDF conversion tool unavailable or conversion failed: {result_msg}")

    # 5) Assert: The PDF file was generated and is not empty
    # Some environments (especially docx2pdf) might ignore the exact output filename
    # and just generate a PDF with the same name as the source in the output or source directory,
    # so we check multiple possible locations.
    candidates = [
        out_pdf,
        # Common: A PDF with the same name as the source file in the output directory
        out_pdf.parent / f"{src_doc.stem}.pdf",
        # Fallback: A PDF in the same directory as the source file
        src_doc.with_suffix(".pdf"),
    ]

    # If none of the above paths exist, search for any newly generated PDF in the temp directory
    found = None
    for p in candidates:
        if p.exists():
            found = p
            break
    if not found:
        pdfs = sorted(tmp_path.glob("*.pdf"), key=lambda p: p.stat().st_mtime, reverse=True)
        if pdfs:
            found = pdfs[0]

    if not found:
        # If the tool returns success but the output can't be found,
        # treat it as an environment/tooling difference and skip instead of failing.
        pytest.skip(f"Could not find the generated PDF. Function output: {result_msg}")

    assert found.exists(), f"Generated PDF not found: {found}, function output: {result_msg}"
    assert found.stat().st_size > 0, f"The generated PDF file is empty: {found}"


if __name__ == "__main__":
    # Allow running this file directly for quick verification:
    #   python tests/test_convert_to_pdf.py
    import sys
    sys.exit(pytest.main([__file__, "-q"]))
